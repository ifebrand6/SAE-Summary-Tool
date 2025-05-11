from docx import Document
import re
from typing import List, Dict, Any, Optional

def extract_table_number_and_title(paragraph_text: str):
    """
    Extracts the table number (e.g., '2.1.1') and title from a paragraph.
    """
    match = re.search(r'Table\s*(\d+(?:\.\d+)+)', paragraph_text)
    table_number = match.group(1) if match else None
    table_title = paragraph_text.strip()
    return table_number, table_title

def find_relevant_tables(doc: Document) -> List[Dict[str, Any]]:
    """
    Finds tables with titles containing 'Summary of Serious Adverse Events'.
    Returns a list of dicts with table_number, table_title, and the table object.
    """
    results = []
    for i, para in enumerate(doc.paragraphs):
        if "Summary of Serious Adverse Events" in para.text:
            table_number, table_title = extract_table_number_and_title(para.text)
           
            for table in doc.tables:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                if set(["Preferred Term", "Placebo", "Compound X"]).issubset(set(headers)):
                    results.append({
                        "table_number": table_number,
                        "table_title": table_title,
                        "table": table,
                        "headers": headers
                    })
                    break  # Only the first relevant table after the paragraph
    return results

def parse_table_data(table, headers) -> (List[Dict[str, str]], Optional[Dict[str, str]]):
    """
    Parses table rows into a list of dicts and finds the total row.
    """
    data = []
    total_row = None
    for row in table.rows[1:]:
        row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
        data.append(row_data)
    for row in data:
        if "total" in row[headers[0]].lower():
            total_row = row
            break
    return data, total_row

def generate_summary(data, total_row) -> List[str]:
    """
    Generates summary sentences for the table data.
    """
    summary = []
    rows = [r for r in data if r != total_row]
    try:
        total_placebo = int(total_row["Placebo"])
        total_compound = int(total_row["Compound X"])
    except (TypeError, ValueError, KeyError):
        total_placebo = total_compound = None
    # Total SAE events
    total_events = sum(
        int(row["Placebo"]) + int(row["Compound X"]) for row in rows
        if row["Placebo"].isdigit() and row["Compound X"].isdigit()
    )
    if total_events > 0:
        summary.append(f"There were {total_events} total serious adverse events reported.")
        for row in rows:
            preferred_term = row["Preferred Term"]
            try:
                placebo_count = int(row["Placebo"])
                compound_count = int(row["Compound X"])
            except (ValueError, KeyError):
                continue
            if total_placebo and total_placebo > 0 and placebo_count > 0:
                placebo_pct = round((placebo_count / total_placebo) * 100, 1)
                summary.append(f"{placebo_pct}% of Placebo participants experienced {preferred_term.lower()}.")
            if total_compound and total_compound > 0 and compound_count > 0:
                compound_pct = round((compound_count / total_compound) * 100, 1)
                summary.append(f"{compound_pct}% of Compound X participants experienced {preferred_term.lower()}.")
    else:
        summary.append("No serious adverse events were reported.")
    return summary

def parse_sae_tables(docx_path: str) -> List[Dict[str, Any]]:
    """
    Main entry point: parses the .docx file and returns structured SAE summary results.
    """
    doc = Document(docx_path)
    results = []
    relevant_tables = find_relevant_tables(doc)
    for tbl in relevant_tables:
        data, total_row = parse_table_data(tbl["table"], tbl["headers"])
        summary = generate_summary(data, total_row)
        results.append({
            "table_number": tbl["table_number"],
            "table_title": tbl["table_title"],
            "summary": summary
        })
    if not results:
        results.append({
            "table_number": None,
            "table_title": None,
            "summary": ["No serious adverse events were reported."]
        })
    return results
