from docx import Document
import re

def parse_sae_tables(docx_path):
    doc = Document(docx_path)
    results = []
    sae_found = False
    for i, para in enumerate(doc.paragraphs):
        if "Summary of Serious Adverse Events" in para.text:
            # Extract table number (e.g., '2.1.1') if present
            match = re.search(r'Table\s*(\d+(?:\.\d+)+)', para.text)
            table_number = match.group(1) if match else None
            table_title = para.text.strip()
            for table in doc.tables:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                if set(["Preferred Term", "Placebo", "Compound X"]).issubset(set(headers)):
                    data = []
                    for row in table.rows[1:]:
                        row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
                        data.append(row_data)
                    total_row = next((r for r in data if "total" in r[headers[0]].lower()), None)
                    rows = [r for r in data if r != total_row]
                    summary = []
                    # Try to get totals
                    try:
                        total_placebo = int(total_row["Placebo"])
                        total_compound = int(total_row["Compound X"])
                    except (TypeError, ValueError, KeyError):
                        total_placebo = total_compound = None
                    # 1. Total SAE events
                    total_events = sum(
                        int(row["Placebo"]) + int(row["Compound X"]) for row in rows
                        if row["Placebo"].isdigit() and row["Compound X"].isdigit()
                    )
                    if total_events > 0:
                        sae_found = True
                        summary.append(f"There were {total_events} total serious adverse events reported.")
                        # 2. Percentages for each preferred term
                        for row in rows:
                            preferred_term = row["Preferred Term"]
                            try:
                                placebo_count = int(row["Placebo"])
                                compound_count = int(row["Compound X"])
                            except (ValueError, KeyError):
                                continue
                            if total_placebo and total_placebo > 0 and placebo_count > 0:
                                placebo_pct = round((placebo_count / total_placebo) * 100, 1)
                                summary.append(f"{placebo_pct}% of Placebo participants experienced {preferred_term.lower()}.")
                            if total_compound and total_compound > 0 and compound_count > 0:
                                compound_pct = round((compound_count / total_compound) * 100, 1)
                                summary.append(f"{compound_pct}% of Compound X participants experienced {preferred_term.lower()}.")
                    else:
                        summary.append("No serious adverse events were reported.")
                    results.append({
                        "table_number": table_number,
                        "table_title": table_title,
                        "summary": summary
                    })
                    break
    if not results:
        results.append({
            "table_number": None,
            "table_title": None,
            "summary": ["No serious adverse events were reported."]
        })
    return results
